import io
import html
import zipfile
from typing import Dict, List, Tuple

import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT

def export_to_pdf(translated_text: str, original_filename: str = "Paper") -> io.BytesIO:
    """
    Genera un archivo PDF académico formateado a partir del texto traducido al español.
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'AcademicTitle',
        parent=styles['Title'],
        fontName='Helvetica-Bold',
        fontSize=18,
        leading=22,
        alignment=TA_CENTER,
        spaceAfter=15
    )

    subtitle_style = ParagraphStyle(
        'AcademicSubtitle',
        parent=styles['Normal'],
        fontName='Helvetica-Oblique',
        fontSize=10,
        leading=13,
        alignment=TA_CENTER,
        spaceAfter=20
    )

    h1_style = ParagraphStyle(
        'AcademicH1',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=14,
        leading=18,
        spaceBefore=14,
        spaceAfter=6
    )

    h2_style = ParagraphStyle(
        'AcademicH2',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=15,
        spaceBefore=10,
        spaceAfter=4
    )

    h3_style = ParagraphStyle(
        'AcademicH3',
        parent=styles['Heading3'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=14,
        spaceBefore=8,
        spaceAfter=3
    )

    body_style = ParagraphStyle(
        'AcademicBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        alignment=TA_JUSTIFY,
        spaceAfter=8
    )

    story = []

    # Título del documento
    clean_title = html.escape(original_filename.replace('_', ' '))
    story.append(Paragraph(f"Traducción al Español: {clean_title}", title_style))
    story.append(Paragraph("Artículo Científico / Traducción Asistida por Gemini", subtitle_style))
    story.append(Spacer(1, 10))

    # Párrafos
    paragraphs = translated_text.split("\n\n")
    for para in paragraphs:
        p_text = para.strip()
        if not p_text:
            continue
        
        # Escapar caracteres especiales para ReportLab XML parser
        safe_text = html.escape(p_text)
        
        # Identificar encabezados Markdown
        if p_text.startswith("# "):
            clean_head = html.escape(p_text.replace("# ", "").strip())
            story.append(Paragraph(clean_head, h1_style))
        elif p_text.startswith("## "):
            clean_head = html.escape(p_text.replace("## ", "").strip())
            story.append(Paragraph(clean_head, h2_style))
        elif p_text.startswith("### "):
            clean_head = html.escape(p_text.replace("### ", "").strip())
            story.append(Paragraph(clean_head, h3_style))
        else:
            # Reemplazar saltos de línea internos por <br/>
            safe_text = safe_text.replace("\n", "<br/>")
            story.append(Paragraph(safe_text, body_style))

    doc.build(story)
    buffer.seek(0)
    return buffer

def export_to_docx(translated_text: str, original_filename: str = "Paper") -> io.BytesIO:
    """Genera un archivo DOCX formateado a partir del texto traducido."""
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    title = doc.add_heading(f"Traducción: {original_filename}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    paragraphs = translated_text.split("\n\n")
    for para in paragraphs:
        p_text = para.strip()
        if not p_text:
            continue
        
        if p_text.startswith("# "):
            doc.add_heading(p_text.replace("# ", ""), level=1)
        elif p_text.startswith("## "):
            doc.add_heading(p_text.replace("## ", ""), level=2)
        elif p_text.startswith("### "):
            doc.add_heading(p_text.replace("### ", ""), level=3)
        else:
            p = doc.add_paragraph(p_text)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(6)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def export_to_markdown(translated_text: str) -> str:
    """Retorna el texto en formato Markdown plano."""
    return translated_text

def export_to_txt(translated_text: str) -> str:
    """Retorna el texto en formato TXT plano."""
    return translated_text

def create_zip_of_pdfs(pdf_files: List[Tuple[str, bytes]]) -> io.BytesIO:
    """
    Crea un archivo ZIP en memoria conteniendo una lista de tuplas (nombre_archivo.pdf, bytes_del_pdf).
    """
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        for filename, data in pdf_files:
            zip_file.writestr(filename, data)
    zip_buffer.seek(0)
    return zip_buffer
